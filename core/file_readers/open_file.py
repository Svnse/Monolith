from __future__ import annotations

import csv
import gzip
import io
import json
import tarfile
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

TEXT_EXTENSIONS = frozenset({
    ".txt", ".md", ".markdown", ".py", ".json", ".jsonl", ".yaml", ".yml",
    ".csv", ".tsv", ".log", ".xml", ".html", ".htm", ".css", ".js", ".jsx",
    ".ts", ".tsx", ".toml", ".ini", ".cfg", ".conf", ".sh", ".bat", ".ps1",
    ".sql", ".r", ".rs", ".go", ".java", ".kt", ".c", ".cpp", ".cc", ".h",
    ".hpp", ".cs", ".php", ".rb", ".swift", ".scala", ".lua",
})
PDF_EXTENSIONS = frozenset({".pdf"})
DOCX_EXTENSIONS = frozenset({".docx"})
XLSX_EXTENSIONS = frozenset({".xlsx", ".xlsm"})
IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"})
ZIP_EXTENSIONS = frozenset({".zip"})
TAR_EXTENSIONS = frozenset({".tar", ".tgz", ".tbz", ".tbz2", ".txz"})
UNSUPPORTED_ARCHIVE_EXTENSIONS = frozenset({".7z", ".rar"})
GZIP_EXTENSIONS = frozenset({".gz"})

MAX_ARCHIVE_ENTRIES = 2000
MAX_ARCHIVE_MEMBER_BYTES = 8 * 1024 * 1024
DEFAULT_MAX_CHARS = 8000


@dataclass(frozen=True)
class OpenFileRequest:
    path: Path
    path_str: str
    max_chars: int = DEFAULT_MAX_CHARS
    offset: int = 0
    member: str | None = None
    sheet: str | None = None
    max_rows: int = 80
    max_members: int = 80


def open_file(
    path: str | Path,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    offset: int = 0,
    member: str | None = None,
    sheet: str | None = None,
    max_rows: int = 80,
    max_members: int = 80,
) -> str:
    raw_path = str(path)
    resolved = Path(path).expanduser().resolve()
    req = OpenFileRequest(
        path=resolved,
        path_str=raw_path,
        max_chars=_clamp_int(max_chars, DEFAULT_MAX_CHARS, 200, 50000),
        offset=_clamp_int(offset, 0, 0, 10**9),
        member=(str(member).strip() or None) if member is not None else None,
        sheet=(str(sheet).strip() or None) if sheet is not None else None,
        max_rows=_clamp_int(max_rows, 80, 1, 1000),
        max_members=_clamp_int(max_members, 80, 1, 500),
    )

    if not req.path.exists():
        return f"[open_file: path does not exist - {req.path}]"
    if not req.path.is_file():
        return f"[open_file: path is not a file - {req.path}]"

    suffix = _normalized_suffix(req.path)
    if suffix in TEXT_EXTENSIONS:
        return _open_text_file(req, kind="text")
    if suffix in ZIP_EXTENSIONS:
        return _open_zip(req)
    if suffix in TAR_EXTENSIONS or _is_tar_gz(req.path):
        return _open_tar(req)
    if suffix in GZIP_EXTENSIONS and not _is_tar_gz(req.path):
        return _open_gzip(req)
    if suffix in PDF_EXTENSIONS:
        return _open_pdf(req)
    if suffix in DOCX_EXTENSIONS:
        return _open_docx(req)
    if suffix in XLSX_EXTENSIONS:
        return _open_xlsx(req)
    if suffix in IMAGE_EXTENSIONS:
        return _open_image(req)
    if suffix in UNSUPPORTED_ARCHIVE_EXTENSIONS:
        return (
            f"[open_file: {req.path.name} ({req.path}) | kind=archive | unsupported]\n"
            f"{suffix} archives need an optional backend. Supported now: .zip, .tar, .tgz, .gz."
        )
    return _open_unknown(req)


def _clamp_int(value: object, default: int, low: int, high: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default
    return max(low, min(parsed, high))


def _normalized_suffix(path: Path) -> str:
    name = path.name.lower()
    if name.endswith(".tar.gz"):
        return ".tgz"
    if name.endswith(".tar.bz2"):
        return ".tbz2"
    if name.endswith(".tar.xz"):
        return ".txz"
    return path.suffix.lower()


def _is_tar_gz(path: Path) -> bool:
    lowered = path.name.lower()
    return lowered.endswith((".tar.gz", ".tar.bz2", ".tar.xz"))


def _safe_stat(path: Path) -> str:
    try:
        size = path.stat().st_size
    except OSError:
        return "size=?"
    return f"size={size:,}B"


def _hint_path(value: str) -> str:
    return value.replace("\\", "/").replace('"', '\\"')


def _hint_value(value: str) -> str:
    return value.replace('"', '\\"')


def _window_text(
    text: str,
    req: OpenFileRequest,
    *,
    tool_path_arg: str | None = None,
    continuation_args: str | None = None,
) -> tuple[str, str]:
    total = len(text)
    start = min(req.offset, total)
    chunk = text[start:start + req.max_chars]
    end = start + len(chunk)
    pct = int((end / total) * 100) if total else 100
    suffix = ""
    if end < total:
        path_arg = tool_path_arg if tool_path_arg is not None else _hint_path(req.path_str)
        args = continuation_args if continuation_args is not None else f'path="{path_arg}"'
        suffix = (
            f"\n[PARTIAL - {total - end:,} chars remaining. "
            f"To continue: open_file({args}, offset={end})]"
        )
    return chunk + suffix, f"chars {start}-{end} of {total} | {pct}% shown"


def _open_text_file(req: OpenFileRequest, *, kind: str) -> str:
    try:
        text = req.path.read_text(encoding="utf-8", errors="replace")
    except PermissionError:
        return f"[open_file: permission denied - {req.path}]"
    except Exception as exc:
        return f"[open_file: error reading text - {exc}]"
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind={kind} | {span}]\n{body}"


def _open_unknown(req: OpenFileRequest) -> str:
    try:
        sample = req.path.read_bytes()[:4096]
    except Exception as exc:
        return f"[open_file: error reading metadata - {exc}]"
    if _looks_like_text(sample):
        return _open_text_file(req, kind="text/unknown-extension")
    return (
        f"[open_file: {req.path.name} ({req.path}) | kind=binary | {_safe_stat(req.path)}]\n"
        "No text extractor is registered for this file type."
    )


def _looks_like_text(data: bytes) -> bool:
    if not data:
        return True
    if b"\x00" in data:
        return False
    try:
        data.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False


def _archive_entry_lines(entries: Iterable[tuple[str, int, bool]], *, max_members: int) -> list[str]:
    lines: list[str] = []
    for index, (name, size, is_dir) in enumerate(entries):
        if index >= max_members:
            lines.append(f"  ... {index} shown; raise max_members to list more")
            break
        suffix = "/" if is_dir and not name.endswith("/") else ""
        size_text = "" if is_dir else f" ({size:,}B)"
        lines.append(f"  {name}{suffix}{size_text}")
    return lines


def _open_zip(req: OpenFileRequest) -> str:
    try:
        with zipfile.ZipFile(req.path, "r") as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ARCHIVE_ENTRIES:
                return (
                    f"[open_file: {req.path.name} ({req.path}) | kind=zip | blocked]\n"
                    f"Archive has {len(infos)} entries; limit is {MAX_ARCHIVE_ENTRIES}."
                )
            if req.member:
                return _open_zip_member(req, archive, infos)
            entries = ((info.filename, int(info.file_size), info.is_dir()) for info in infos)
            lines = _archive_entry_lines(entries, max_members=req.max_members)
            hint = (
                f'To preview a text member: open_file(path="{req.path_str}", '
                'member="path/in/archive.txt").'
            )
            return (
                f"[open_file: {req.path.name} ({req.path}) | kind=zip | "
                f"entries={len(infos)} | {_safe_stat(req.path)}]\n"
                + "\n".join(lines)
                + f"\n{hint}"
            )
    except zipfile.BadZipFile:
        return f"[open_file: invalid zip archive - {req.path}]"
    except Exception as exc:
        return f"[open_file: error reading zip - {exc}]"


def _open_zip_member(req: OpenFileRequest, archive: zipfile.ZipFile, infos: list[zipfile.ZipInfo]) -> str:
    lookup = {info.filename: info for info in infos}
    info = lookup.get(req.member or "")
    if info is None:
        return f"[open_file: zip member not found - {req.member}]"
    if info.is_dir():
        return f"[open_file: zip member is a directory - {req.member}]"
    if info.file_size > MAX_ARCHIVE_MEMBER_BYTES:
        return (
            f"[open_file: zip member too large - {req.member} "
            f"({info.file_size:,}B; limit {MAX_ARCHIVE_MEMBER_BYTES:,}B)]"
        )
    data = archive.read(info)
    if not _looks_like_text(data[:4096]):
        return f"[open_file: zip member is binary - {req.member} ({info.file_size:,}B)]"
    text = data.decode("utf-8", errors="replace")
    continuation_args = (
        f'path="{_hint_path(req.path_str)}", member="{_hint_value(req.member or "")}"'
    )
    body, span = _window_text(text, req, continuation_args=continuation_args)
    return (
        f"[open_file: {req.path.name}!{req.member} ({req.path}) | kind=zip-member | {span}]\n"
        f"{body}"
    )


def _open_tar(req: OpenFileRequest) -> str:
    try:
        with tarfile.open(req.path, "r:*") as archive:
            members = archive.getmembers()
            if len(members) > MAX_ARCHIVE_ENTRIES:
                return (
                    f"[open_file: {req.path.name} ({req.path}) | kind=tar | blocked]\n"
                    f"Archive has {len(members)} entries; limit is {MAX_ARCHIVE_ENTRIES}."
                )
            if req.member:
                return _open_tar_member(req, archive, members)
            entries = ((m.name, int(m.size), m.isdir()) for m in members)
            lines = _archive_entry_lines(entries, max_members=req.max_members)
            hint = (
                f'To preview a text member: open_file(path="{req.path_str}", '
                'member="path/in/archive.txt").'
            )
            return (
                f"[open_file: {req.path.name} ({req.path}) | kind=tar | "
                f"entries={len(members)} | {_safe_stat(req.path)}]\n"
                + "\n".join(lines)
                + f"\n{hint}"
            )
    except tarfile.TarError:
        return f"[open_file: invalid tar archive - {req.path}]"
    except Exception as exc:
        return f"[open_file: error reading tar - {exc}]"


def _open_tar_member(req: OpenFileRequest, archive: tarfile.TarFile, members: list[tarfile.TarInfo]) -> str:
    member = next((m for m in members if m.name == req.member), None)
    if member is None:
        return f"[open_file: tar member not found - {req.member}]"
    if member.isdir():
        return f"[open_file: tar member is a directory - {req.member}]"
    if member.size > MAX_ARCHIVE_MEMBER_BYTES:
        return (
            f"[open_file: tar member too large - {req.member} "
            f"({member.size:,}B; limit {MAX_ARCHIVE_MEMBER_BYTES:,}B)]"
        )
    stream = archive.extractfile(member)
    if stream is None:
        return f"[open_file: tar member cannot be read - {req.member}]"
    data = stream.read(MAX_ARCHIVE_MEMBER_BYTES + 1)
    if not _looks_like_text(data[:4096]):
        return f"[open_file: tar member is binary - {req.member} ({member.size:,}B)]"
    text = data.decode("utf-8", errors="replace")
    continuation_args = (
        f'path="{_hint_path(req.path_str)}", member="{_hint_value(req.member or "")}"'
    )
    body, span = _window_text(text, req, continuation_args=continuation_args)
    return (
        f"[open_file: {req.path.name}!{req.member} ({req.path}) | kind=tar-member | {span}]\n"
        f"{body}"
    )


def _open_gzip(req: OpenFileRequest) -> str:
    try:
        with gzip.open(req.path, "rb") as fh:
            data = fh.read(MAX_ARCHIVE_MEMBER_BYTES + 1)
    except OSError:
        return f"[open_file: invalid gzip file - {req.path}]"
    except Exception as exc:
        return f"[open_file: error reading gzip - {exc}]"
    if len(data) > MAX_ARCHIVE_MEMBER_BYTES:
        return (
            f"[open_file: gzip payload too large - {req.path.name} "
            f"(limit {MAX_ARCHIVE_MEMBER_BYTES:,}B)]"
        )
    if not _looks_like_text(data[:4096]):
        return f"[open_file: gzip payload is binary - {req.path.name}]"
    text = data.decode("utf-8", errors="replace")
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind=gzip-text | {span}]\n{body}"


def _open_pdf(req: OpenFileRequest) -> str:
    try:
        import fitz  # type: ignore
    except Exception:
        return _open_pdf_pypdf(req)
    try:
        doc = fitz.open(str(req.path))
        page_texts = []
        for idx, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            if text.strip():
                page_texts.append(f"\n\n[page {idx}]\n{text.strip()}")
        doc.close()
    except Exception as exc:
        return f"[open_file: error reading pdf - {exc}]"
    text = "".join(page_texts).strip()
    if not text:
        return (
            f"[open_file: {req.path.name} ({req.path}) | kind=pdf | no-text-layer]\n"
            "No selectable PDF text was found. OCR backend is needed for scanned/image-only pages."
        )
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind=pdf-text | {span}]\n{body}"


def _open_pdf_pypdf(req: OpenFileRequest) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            return (
                f"[open_file: {req.path.name} ({req.path}) | kind=pdf | dependency-missing]\n"
                "PDF text extraction needs PyMuPDF or pypdf. "
                "If this is a scanned PDF, OCR backend is still required after rendering pages."
            )
    try:
        reader = PdfReader(str(req.path))
        page_texts = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                page_texts.append(f"\n\n[page {idx}]\n{text.strip()}")
    except Exception as exc:
        return f"[open_file: error reading pdf - {exc}]"
    text = "".join(page_texts).strip()
    if not text:
        return (
            f"[open_file: {req.path.name} ({req.path}) | kind=pdf | no-text-layer]\n"
            "No selectable PDF text was found. OCR backend is needed for scanned/image-only pages."
        )
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind=pdf-text | {span}]\n{body}"


def _open_docx(req: OpenFileRequest) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return _open_docx_zip(req)
    try:
        doc = Document(str(req.path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
    except Exception as exc:
        return f"[open_file: error reading docx - {exc}]"
    text = "\n".join(parts).strip()
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind=docx | {span}]\n{body}"


def _open_docx_zip(req: OpenFileRequest) -> str:
    try:
        with zipfile.ZipFile(req.path, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                return f"[open_file: invalid docx archive - missing word/document.xml]"
            xml_bytes = archive.read("word/document.xml")
    except zipfile.BadZipFile:
        return f"[open_file: invalid docx archive - {req.path}]"
    except Exception as exc:
        return f"[open_file: error reading docx - {exc}]"
    try:
        root = ET.fromstring(xml_bytes)
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []
        for paragraph in root.iter(f"{w_ns}p"):
            text = "".join(node.text or "" for node in paragraph.iter(f"{w_ns}t")).strip()
            if text:
                parts.append(text)
    except Exception as exc:
        return f"[open_file: error parsing docx - {exc}]"
    text = "\n".join(parts).strip()
    body, span = _window_text(text, req)
    return f"[open_file: {req.path.name} ({req.path}) | kind=docx | {span}]\n{body}"


def _open_xlsx(req: OpenFileRequest) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception:
        return (
            f"[open_file: {req.path.name} ({req.path}) | kind=xlsx | dependency-missing]\n"
            "XLSX extraction needs openpyxl (`pip install openpyxl`)."
        )
    try:
        wb = openpyxl.load_workbook(req.path, read_only=True, data_only=True)
        ws = wb[req.sheet] if req.sheet and req.sheet in wb.sheetnames else wb[wb.sheetnames[0]]
        output = io.StringIO()
        writer = csv.writer(output, delimiter="\t", lineterminator="\n")
        for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if idx > req.max_rows:
                break
            writer.writerow(["" if value is None else str(value) for value in row])
        sheet_title = ws.title
        sheet_names = ", ".join(wb.sheetnames)
        wb.close()
    except Exception as exc:
        return f"[open_file: error reading xlsx - {exc}]"
    text = output.getvalue().rstrip()
    body, span = _window_text(text, req)
    return (
        f"[open_file: {req.path.name} ({req.path}) | kind=xlsx | sheet={sheet_title} | "
        f"sheets={sheet_names} | rows_shown<={req.max_rows} | {span}]\n{body}"
    )


def _open_image(req: OpenFileRequest) -> str:
    dims = ""
    try:
        from PIL import Image  # type: ignore

        with Image.open(req.path) as image:
            dims = f" | dimensions={image.width}x{image.height}"
    except Exception:
        dims = ""
    return (
        f"[open_file: {req.path.name} ({req.path}) | kind=image{dims} | {_safe_stat(req.path)}]\n"
        "Image text requires an OCR backend. No OCR backend is configured yet."
    )
